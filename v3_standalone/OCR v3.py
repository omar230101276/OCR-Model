
import io
import os
import sys
import argparse
import easyocr
import cv2
import numpy as np
import pandas as pd
import re
import shutil
import fitz # PyMuPDF
import docx
from PIL import Image

# -----------------------------------------------------------------------------
# 1. OCR ENGINE (from src/core_ocr.py)
# -----------------------------------------------------------------------------
class OCREngine:
    def __init__(self, languages=['en', 'ar'], gpu=True):
        """
        Initialize the OCR engine.
        :param languages: List of supported languages (default: English and Arabic)
        :param gpu: Use GPU for acceleration
        """
        # Determine model storage path
        # Try to find 'models_cache' in current dir or parent dir to avoid redownloading
        cwd = os.getcwd()
        possible_paths = [
            os.path.join(cwd, 'models_cache'),
            os.path.join(os.path.dirname(cwd), 'models_cache'), # One level up
            os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'models_cache') # Relative to script
        ]
        
        model_dir = possible_paths[0] # Default
        for path in possible_paths:
            if os.path.exists(path):
                model_dir = path
                break
        
        if not os.path.exists(model_dir):
            os.makedirs(model_dir)

        print(f"Loading OCR model for languages: {languages} (GPU={gpu})...")
        print(f"Models will be stored in: {model_dir}")
        
        self.reader = easyocr.Reader(languages, gpu=gpu, model_storage_directory=model_dir, download_enabled=True)

    def read_image(self, image_path, detail=1):
        """
        Read text from an image or PDF.
        :param image_path: Path to the image or PDF
        :param detail: Detail level (1 for boxes and text, 0 for text only)
        :return: Reading results
        """
        if image_path.lower().endswith('.pdf'):
            print(f"Detected PDF: {image_path}. Converting to images...")
            images = convert_pdf_to_images(image_path)
            all_results = []
            for i, img in enumerate(images):
                print(f"Processing page {i+1}/{len(images)}...")
                results = self.reader.readtext(img, detail=detail)
                all_results.extend(results)
                all_results.extend(results)
            return all_results
            
        elif image_path.lower().endswith('.docx'):
            print(f"Detected DOCX: {image_path}. extracting text and images...")
            return process_docx(image_path, self, detail=detail)

        return self.reader.readtext(image_path, detail=detail)

    def read_image_from_array(self, image_array, detail=1):
        """
        Read text from an image array (NumPy array).
        Useful when cropping images or processing before reading.
        """
        return self.reader.readtext(image_array, detail=detail)

# -----------------------------------------------------------------------------
# 1.5 PDF UTILS (Inline for standalone)
# -----------------------------------------------------------------------------
def convert_pdf_to_images(pdf_path, zoom=2.0):
    """
    Convert a PDF file to a list of images (numpy arrays).
    """
    images = []
    try:
        doc = fitz.open(pdf_path)
        mat = fitz.Matrix(zoom, zoom)
        
        for page in doc:
            pix = page.get_pixmap(matrix=mat)
            img_array = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.h, pix.w, pix.n)
            if pix.n == 4:
                img_array = img_array[..., :3]
            images.append(img_array)
            
        doc.close()
    except Exception as e:
        print(f"Error converting PDF to images: {e}")
        return []
    return images


# -----------------------------------------------------------------------------
# 1.6 DOCX UTILS (Inline for standalone)
# -----------------------------------------------------------------------------
def process_docx(docx_path, ocr_engine, detail=1):
    """
    Extract text and images from a DOCX file.
    """
    results = []
    
    try:
        doc = docx.Document(docx_path)
        
        # 1. Extract Text from Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                results.append(([[0,0], [1,0], [1,1], [0,1]], para.text.strip(), 1.0))
                
        # 2. Extract Text from Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                         results.append(([[0,0], [1,0], [1,1], [0,1]], cell.text.strip(), 1.0))

        # 3. Extract Images (Advanced)
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    image_data = rel.target_part.blob
                    pil_img = Image.open(io.BytesIO(image_data))
                    pil_img = pil_img.convert('RGB')
                    img_array = np.array(pil_img)
                    
                    print(f"Found embedded image of size {img_array.shape}, running OCR...")
                    ocr_results = ocr_engine.read_image_from_array(img_array, detail=detail)
                    results.extend(ocr_results)
                except Exception as img_e:
                    print(f"Failed to process an embedded image: {img_e}")

    except Exception as e:
        print(f"Error processing DOCX: {e}")
        return []

    # Filter results based on detail level
    if detail == 0:
        return [r[1] for r in results]
    else:
        return results

# -----------------------------------------------------------------------------
# -----------------------------------------------------------------------------
# 2. SPECIFICATION EXTRACTOR (from src/extraction.py)
# -----------------------------------------------------------------------------
class SpecificationExtractor:
    def __init__(self):
        # Using broad patterns from original ocr.py to catch OCR errors
        self.patterns = {
            "cable_type": {
                # Copper, C0pper, COpp er, etc.
                "en": r"\b(C[o0]pp[\s]*[e3]r|Cu|Aluminium|Aluminum|Al)\b\s*(?:C[@a]ble|Conductor)?",
                "ar": r"(نحاس|الومنيوم)"
            },
            "voltage": {
                # 450/750V, 450 / 7S0 V, 0.6/1kV
                "en": r"(\d[\d\s\.]*[/]?[\d\sS]*\s*[kK]?[vV])", 
                "ar": r"(\d+\s*פولت|جهد\s*\d+)"
            },
            "current_rating": {
                 # 32 A, 3 2 A
                 "en": r"(\d[\d\s]*\s*(?:Amps?|A)\b)",
                 "ar": r"(\d+\s*امبير)"
            },
            "insulation": {
                "en": r"(XLPE|PVC)",
                "ar": r"(اكس ال بي اي|بي في سي)"
            },
            "conductor_count": {
                "en": r"(\d+)\s*(?:Core|Cores|x)",
                "ar": r"(\d+)\s*(?:كور|خط)"
            },
            "conductor_size": { 
                # 6 mm2, 6m m 2, 6  mm2
                "en": r"(\d+(?:[\s]*[xX][\s]*\d+)?[\s]*m[\s]*m[h]?[\s]*[2²\?]?)", 
                "ar": r"(\d+\s*x\s*\d+\s*مم2)"
            },
            "sheath": {
                "en": r"(PVC|HDPE|LDPE|Lead|LAZH|LSOH|MDPE|EPR|PUR|TPU|Neoprene|Rubber|LSZH)\s*(?:Sheath|Jacket)\s*(?:Sheath|Jacket)?",
                "ar": r"(غلاف\s*بي في سي)"
            },
            "operating_temperature": {
                # 40 C, 4O C
                "en": r"(\d+[O0]?[\s]*(?:°|\*|deg|degrees)?[\s]*C)",
                "ar": r"(\d+\s*درجة)"
            },
            "insulation_resistance": {
                # 20 MO.km, 20M O.km
                "en": r"(\d+[\s]*M?[O\u03a9][\s]*[\.]?k?m)",
                "ar": r"(\d+\s*ميج اوم)"
            },
            "armor": {
                # Steel Wire Armor, Steel Tape Armor, SWA, STA (word boundaries)
                "en": r"(Stee[l1][\s]*[WT][l1Iae3pi]+[\s]*Armo[r0x]|\bSWA\b|\bSTA\b|SWA|AWA|ATA|GSWA|GSTA|CWA|BWA)",
                "ar": r"(تسليح|مسلح)"
            }
        }

    def extract_specs(self, text):
        """
        Extract specifications from full text.
        """
        specs = {}
        for key, lang_patterns in self.patterns.items():
            specs[key] = None
            
            # Search English (Prioritize original logic)
            match_en = re.search(lang_patterns["en"], text, re.IGNORECASE)
            if match_en:
                # If we have groups, use group 1, unless it's the whole match we want
                if len(match_en.groups()) > 0:
                     specs[key] = match_en.group(1)
                else:
                     specs[key] = match_en.group(0)
                continue
            
            # Search Arabic
            match_ar = re.search(lang_patterns["ar"], text, re.IGNORECASE)
            if match_ar:
                specs[key] = match_ar.group(0)
        
        return self.clean_specs(specs)

    def clean_specs(self, specs):
        """
        Clean and correct extracted data using original logic.
        """
        # voltage correction logic
        if specs.get("voltage"):
            val = specs["voltage"]
            # Remove spaces
            val = val.replace(" ", "")
            # Fix S -> 5
            val = val.replace("S", "5").replace("s", "5")
            
            if "6" in val and "v" in val.lower() and ("1000" in val or "1k" in val):
                 specs["voltage"] = "600/1000V"
            # Attempt to normalize 450/750
            if "450" in val and "750" in val:
                specs["voltage"] = "450/750V"
            else:
                specs["voltage"] = val
        
        # conductor/size correction
        if specs.get("conductor_size"):
            val = specs["conductor_size"]
            val = val.replace(" ", "")
            val = val.replace("mh", "mm").replace("?", "2")
            if "mm" in val and not val.endswith("2"): # Append 2 if missing (e.g. 6mm -> 6mm2)
                 val += "2"
            specs["conductor_size"] = val

        # current rating correction
        if specs.get("current_rating"):
            val = specs["current_rating"]
            val = val.replace(" ", "")
            specs["current_rating"] = val

        # armor correction
        if specs.get("armor"):
             val = specs["armor"]
             import re
             if re.search(r"Stee[l1]", val, re.IGNORECASE):
                 specs["armor"] = "Steel Wire Armor"
             else:
                 specs["armor"] = val.replace("Armox", "Armor").replace("armox", "armor")
        
        # cable type cleanup
        if specs.get("cable_type"):
            val = specs["cable_type"].lower()
            if "c" in val and ("p" in val or "o" in val) and "r" in val: # Loose check for copper
                specs["cable_type"] = "Copper"
            elif "al" in val: 
                specs["cable_type"] = "Aluminum"

        # sheath cleanup
        if specs.get("sheath"):
             val = specs["sheath"].upper()
             if "PVC" in val: specs["sheath"] = "PVC"
             elif "HDPE" in val: specs["sheath"] = "HDPE"

        # resistance cleanup (OCR often reads Omega as O or 0)
        if specs.get("insulation_resistance"):
             specs["insulation_resistance"] = specs["insulation_resistance"].replace("O", "Ω").replace(" ", "")

        return specs

# -----------------------------------------------------------------------------
# 3. TABLE EXTRACTOR (from src/table_engine.py)
# -----------------------------------------------------------------------------
class TableExtractor:
    def __init__(self, ocr_engine):
        self.ocr = ocr_engine

    def extract_table(self, image_path):
        """
        Attempt to extract a table from the image.
        """
        img = cv2.imread(image_path)
        if img is None:
            raise ValueError("Image not found")

        # 1. Convert to grayscale and process
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        thresh = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)
        thresh = 255 - thresh # Invert colors

        # 2. Detect horizontal and vertical lines
        rows = gray.shape[0]
        vertical_size = rows // 30
        vertical_structure = cv2.getStructuringElement(cv2.MORPH_RECT, (1, vertical_size))
        vertical = cv2.erode(thresh, vertical_structure)
        vertical = cv2.dilate(vertical, vertical_structure)

        cols = gray.shape[1]
        horizontal_size = cols // 30
        horizontal_structure = cv2.getStructuringElement(cv2.MORPH_RECT, (horizontal_size, 1))
        horizontal = cv2.erode(thresh, horizontal_structure)
        horizontal = cv2.dilate(horizontal, horizontal_structure)

        # 3. Combine lines to get the grid
        grid = cv2.add(horizontal, vertical)
        
        # 4. Find cells (Contours)
        contours, _ = cv2.findContours(grid, cv2.RETR_TREE, cv2.CHAIN_APPROX_SIMPLE)

        # Filter out very small contours
        cells = []
        for cnt in contours:
            x, y, w, h = cv2.boundingRect(cnt)
            if w > 20 and h > 10: # Ignore noise
                cells.append((x, y, w, h))

        # 5. Sort cells (Top-to-bottom, then Left-to-right)
        # This sort is simple and might need improvement for complex tables
        cells.sort(key=lambda b: (b[1] // 10, b[0])) # Cluster by rows approximately

        data = []
        current_row_y = -1
        row_data = []

        # Read each cell
        for (x, y, w, h) in cells:
            # Crop cell
            roi = img[y:y+h, x:x+w]
            
            # Read text inside cell
            results = self.ocr.read_image_from_array(roi, detail=0)
            text = " ".join(results).strip()
            
            # Simple logic to determine rows (if Y changes significantly, start new row)
            if current_row_y == -1:
                current_row_y = y
            
            if abs(y - current_row_y) > 20:
                data.append(row_data)
                row_data = []
                current_row_y = y
            
            row_data.append(text)
        
        if row_data:
            data.append(row_data)

        # Convert to DataFrame
        df = pd.DataFrame(data)
        return df

# -----------------------------------------------------------------------------
# 4. MAIN EXECUTION (from main.py)
# -----------------------------------------------------------------------------
def main():
    parser = argparse.ArgumentParser(description="Cable Specification OCR System (Single File Version)")
    parser.add_argument("--image", required=True, help="Path to input image")
    parser.add_argument("--mode", choices=["text", "table"], default="text", help="Operation mode: 'text' for specs/text extraction, 'table' for table extraction")
    parser.add_argument("--output", help="Output path (for CSV tables)")
    parser.add_argument("--langs", default="en", help="Comma-separated list of languages (e.g., 'en,ar')")

    args = parser.parse_args()

    if not os.path.exists(args.image):
        print(f"Error: File {args.image} not found.")
        return

    # Initialize Engine
    languages = args.langs.split(',')
    ocr = OCREngine(languages=languages)

    if args.mode == "text":
        print(f"Reading text from: {args.image} ...")
        results = ocr.read_image(args.image, detail=0)
        full_text = " ".join(results)
        print("\n--- Extracted Text ---")
        print(full_text)
        
        print("\n--- Extracted Specifications ---")
        extractor = SpecificationExtractor()
        specs = extractor.extract_specs(full_text)
        for k, v in specs.items():
            if v:
                print(f"{k}: {v}")
            else:
                print(f"{k}: Not Found")

    elif args.mode == "table":
        print(f"Extracting table from: {args.image} ...")
        table_engine = TableExtractor(ocr)
        try:
            df = table_engine.extract_table(args.image)
            print("\n--- Extracted Table ---")
            print(df)
            
            if args.output:
                df.to_csv(args.output, index=False, encoding='utf-8-sig')
                print(f"Table saved to: {args.output}")
        except Exception as e:
            print(f"Error extracting table: {e}")

if __name__ == "__main__":
    main()
